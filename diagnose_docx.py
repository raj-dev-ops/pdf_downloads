#!/usr/bin/env python3
"""
Diagnostic script to find all images in a DOCX file
"""

import sys
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

def diagnose_docx(docx_path):
    """Diagnose where images are stored in the document"""

    doc = Document(docx_path)

    print(f"Analyzing: {docx_path}\n")
    print("=" * 70)

    # Check document relationships for images
    print("\n1. Checking document relationships for images:")
    print("-" * 70)
    image_parts = []
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            print(f"  Found image relationship: {rel_id} -> {rel.target_ref}")
            image_parts.append((rel_id, rel))

    if not image_parts:
        print("  No image relationships found in document.part.rels")
    else:
        print(f"\n  Total image relationships: {len(image_parts)}")

    # Check inline shapes in paragraphs
    print("\n2. Checking inline shapes in paragraph runs:")
    print("-" * 70)
    inline_count = 0
    for para_idx, paragraph in enumerate(doc.paragraphs):
        for run in paragraph.runs:
            # Check for inline shapes using pic:pic
            pics = run._element.xpath('.//pic:pic')
            if pics:
                inline_count += len(pics)
                print(f"  Paragraph {para_idx}: Found {len(pics)} inline pic(s)")

            # Check for blip elements (actual image references)
            blips = run._element.xpath('.//a:blip')
            if blips:
                for blip in blips:
                    embed_id = blip.get(qn('r:embed'))
                    print(f"    Blip embed ID: {embed_id}")

    if inline_count == 0:
        print("  No inline pictures found in paragraph runs")
    else:
        print(f"\n  Total inline pictures: {inline_count}")

    # Check for drawing objects
    print("\n3. Checking for drawing objects (floating images):")
    print("-" * 70)
    drawing_count = 0
    for para_idx, paragraph in enumerate(doc.paragraphs):
        drawings = paragraph._element.xpath('.//w:drawing')
        if drawings:
            print(f"  Paragraph {para_idx}: Found {len(drawings)} drawing object(s)")
            drawing_count += len(drawings)

            # Look for anchors (floating images)
            anchors = paragraph._element.xpath('.//wp:anchor')
            if anchors:
                print(f"    Contains {len(anchors)} floating/anchored image(s)")

            # Look for inline elements
            inlines = paragraph._element.xpath('.//wp:inline')
            if inlines:
                print(f"    Contains {len(inlines)} inline image(s)")

    if drawing_count == 0:
        print("  No drawing objects found")
    else:
        print(f"\n  Total drawing objects: {drawing_count}")

    # Check in tables
    print("\n4. Checking for images in tables:")
    print("-" * 70)
    table_image_count = 0
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    drawings = paragraph._element.xpath('.//w:drawing')
                    pics = paragraph._element.xpath('.//pic:pic')
                    if drawings or pics:
                        table_image_count += len(drawings) + len(pics)
                        print(f"  Table {table_idx}, Row {row_idx}, Cell {cell_idx}: Found {len(drawings)} drawings, {len(pics)} pics")

    if table_image_count == 0:
        print("  No images found in tables")
    else:
        print(f"\n  Total images in tables: {table_image_count}")

    # Check headers and footers
    print("\n5. Checking headers and footers:")
    print("-" * 70)
    header_footer_count = 0
    for section in doc.sections:
        # Check header
        if section.header:
            for paragraph in section.header.paragraphs:
                drawings = paragraph._element.xpath('.//w:drawing')
                if drawings:
                    header_footer_count += len(drawings)
                    print(f"  Found {len(drawings)} image(s) in header")

        # Check footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                drawings = paragraph._element.xpath('.//w:drawing')
                if drawings:
                    header_footer_count += len(drawings)
                    print(f"  Found {len(drawings)} image(s) in footer")

    if header_footer_count == 0:
        print("  No images found in headers/footers")

    # Summary
    print("\n" + "=" * 70)
    print("SUMMARY:")
    print("=" * 70)
    print(f"  Image relationships in document: {len(image_parts)}")
    print(f"  Inline pictures in paragraphs: {inline_count}")
    print(f"  Drawing objects: {drawing_count}")
    print(f"  Images in tables: {table_image_count}")
    print(f"  Images in headers/footers: {header_footer_count}")
    print("=" * 70)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python diagnose_docx.py <path_to_docx>")
        sys.exit(1)

    diagnose_docx(sys.argv[1])
