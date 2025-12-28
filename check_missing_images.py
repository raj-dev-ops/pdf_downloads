#!/usr/bin/env python3
"""
Check where rId20 and rId7 are located (tables, headers, footers, etc.)
"""

import sys
from docx import Document

def check_missing_images(docx_path):
    """Find where specific images are located"""

    doc = Document(docx_path)

    missing_ids = ['rId20', 'rId7']

    print(f"Searching for: {missing_ids}\n")
    print("=" * 70)

    # Check tables
    print("\nChecking tables:")
    print("=" * 70)
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    # Check for images in this paragraph
                    for run in para.runs:
                        if run._element.xpath('.//pic:pic'):
                            inline_shapes = run._element.xpath('.//a:blip/@r:embed')
                            for rel_id in inline_shapes:
                                if rel_id in missing_ids:
                                    print(f"  FOUND {rel_id} in Table {table_idx}, Row {row_idx}, Cell {cell_idx}")
                                    print(f"    Caption: {para.text[:80]}")

                    # Check VML format
                    pict_elements = para._element.xpath('.//w:pict')
                    if pict_elements:
                        for pict in pict_elements:
                            all_elements = pict.xpath('.//*')
                            for elem in all_elements:
                                for attr_name, attr_value in elem.attrib.items():
                                    if attr_value and isinstance(attr_value, str) and attr_value in missing_ids:
                                        print(f"  FOUND {attr_value} in Table {table_idx}, Row {row_idx}, Cell {cell_idx}")
                                        print(f"    Caption: {para.text[:80]}")

    # Check sections (headers/footers)
    print("\nChecking headers and footers:")
    print("=" * 70)
    for section_idx, section in enumerate(doc.sections):
        # Check header
        try:
            header = section.header
            for para_idx, para in enumerate(header.paragraphs):
                for run in para.runs:
                    if run._element.xpath('.//pic:pic'):
                        inline_shapes = run._element.xpath('.//a:blip/@r:embed')
                        for rel_id in inline_shapes:
                            if rel_id in missing_ids:
                                print(f"  FOUND {rel_id} in Section {section_idx} Header, Para {para_idx}")
                                print(f"    Text: {para.text[:80]}")

                # Check VML
                pict_elements = para._element.xpath('.//w:pict')
                if pict_elements:
                    for pict in pict_elements:
                        all_elements = pict.xpath('.//*')
                        for elem in all_elements:
                            for attr_name, attr_value in elem.attrib.items():
                                if attr_value and isinstance(attr_value, str) and attr_value in missing_ids:
                                    print(f"  FOUND {attr_value} in Section {section_idx} Header, Para {para_idx}")
                                    print(f"    Text: {para.text[:80]}")
        except Exception as e:
            pass

        # Check footer
        try:
            footer = section.footer
            for para_idx, para in enumerate(footer.paragraphs):
                for run in para.runs:
                    if run._element.xpath('.//pic:pic'):
                        inline_shapes = run._element.xpath('.//a:blip/@r:embed')
                        for rel_id in inline_shapes:
                            if rel_id in missing_ids:
                                print(f"  FOUND {rel_id} in Section {section_idx} Footer, Para {para_idx}")
                                print(f"    Text: {para.text[:80]}")

                # Check VML
                pict_elements = para._element.xpath('.//w:pict')
                if pict_elements:
                    for pict in pict_elements:
                        all_elements = pict.xpath('.//*')
                        for elem in all_elements:
                            for attr_name, attr_value in elem.attrib.items():
                                if attr_value and isinstance(attr_value, str) and attr_value in missing_ids:
                                    print(f"  FOUND {attr_value} in Section {section_idx} Footer, Para {para_idx}")
                                    print(f"    Text: {para.text[:80]}")
        except Exception as e:
            pass

    # Also check if these images are in the raw XML but not accessible via python-docx
    print("\nChecking document XML directly:")
    print("=" * 70)
    import zipfile
    from lxml import etree

    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        # Read document.xml
        with zip_ref.open('word/document.xml') as f:
            xml_content = f.read().decode('utf-8')

            for rel_id in missing_ids:
                if rel_id in xml_content:
                    print(f"  {rel_id} IS in document.xml")
                    # Find context
                    import re
                    # Find 100 characters before and after
                    matches = re.finditer(rel_id, xml_content)
                    for match in matches:
                        start = max(0, match.start() - 150)
                        end = min(len(xml_content), match.end() + 150)
                        context = xml_content[start:end]
                        print(f"    Context: ...{context}...")
                else:
                    print(f"  {rel_id} NOT in document.xml - orphaned image in media folder")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python check_missing_images.py <path_to_docx>")
        sys.exit(1)

    check_missing_images(sys.argv[1])
