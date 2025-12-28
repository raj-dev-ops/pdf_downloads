#!/usr/bin/env python3
"""
Find OLE objects (embedded Excel charts) in document order
"""

import sys
from docx import Document

def find_ole_objects(docx_path):
    """Find OLE embedded objects with their positions"""

    doc = Document(docx_path)

    print("Searching for OLE objects (embedded Excel charts) in document order:")
    print("=" * 70)

    found_count = 0

    for para_idx, paragraph in enumerate(doc.paragraphs):
        # Check for w:object elements (embedded objects like Excel charts)
        object_elements = paragraph._element.xpath('.//w:object')

        if object_elements:
            for obj in object_elements:
                found_count += 1

                # Get the relationship ID from v:imagedata within the object
                all_elements = obj.xpath('.//*')
                rel_ids = []

                for elem in all_elements:
                    for attr_name, attr_value in elem.attrib.items():
                        if attr_value and isinstance(attr_value, str) and attr_value.startswith('rId'):
                            rel_ids.append(attr_value)

                # Get nearby text for context (caption)
                caption = paragraph.text.strip()

                # If empty, search nearby paragraphs
                if not caption:
                    for offset in range(1, 4):
                        # Check after
                        if para_idx + offset < len(doc.paragraphs):
                            next_text = doc.paragraphs[para_idx + offset].text.strip()
                            if next_text and ('figure' in next_text.lower() or 'scheme' in next_text.lower()):
                                caption = next_text
                                break

                        # Check before
                        if para_idx - offset >= 0:
                            prev_text = doc.paragraphs[para_idx - offset].text.strip()
                            if prev_text and ('figure' in prev_text.lower() or 'scheme' in prev_text.lower()):
                                caption = prev_text
                                break

                print(f"\nOLE Object #{found_count} at paragraph {para_idx}:")
                print(f"  Relationship IDs: {', '.join(set(rel_ids))}")
                print(f"  Caption: {caption[:100] if caption else '(no caption found)'}")

                # Try to get the image
                for rel_id in set(rel_ids):
                    try:
                        if rel_id in doc.part.related_parts:
                            image_part = doc.part.related_parts[rel_id]
                            print(f"    {rel_id} -> {image_part.partname}")
                    except:
                        pass

    print("\n" + "=" * 70)
    print(f"Total OLE objects found: {found_count}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python find_ole_objects.py <path_to_docx>")
        sys.exit(1)

    find_ole_objects(sys.argv[1])
