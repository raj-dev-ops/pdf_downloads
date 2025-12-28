#!/usr/bin/env python3
"""
Word Document Image Extractor to GIF

Extracts images from Word documents (.doc/.docx) and converts them to GIF format
with specific naming conventions based on image type (Figure or Scheme).
"""

import os
import re
import sys
import argparse
from pathlib import Path
from io import BytesIO
from typing import Tuple, List, Dict

from docx import Document
from docx.oxml.ns import qn
from PIL import Image
import platform

# Optional imports for .doc file support
try:
    import doc2docx
    HAS_DOC2DOCX = True
except ImportError:
    HAS_DOC2DOCX = False

# Try to import win32com for Windows .doc conversion
try:
    import win32com.client
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False


def extract_document_basename(filepath: str) -> str:
    """
    Extract base name from document filename.

    Example:
        IJRIT-11-139-Figures.doc → IJRIT-11-139

    Args:
        filepath: Path to the document file

    Returns:
        Base name without suffixes like -Figures, -Images, -Schemes
    """
    filename = Path(filepath).stem  # Get filename without extension

    # Remove common suffixes
    suffixes_to_remove = ['-Figures', '-figures', '-Images', '-images',
                          '-Schemes', '-schemes', '-Tables', '-tables']

    for suffix in suffixes_to_remove:
        if filename.endswith(suffix):
            filename = filename[:-len(suffix)]
            break

    return filename


def classify_image(caption_text: str) -> str:
    """
    Classify image based on caption text.

    Args:
        caption_text: The caption or alt text from the image

    Returns:
        'g' for figures, 's' for schemes
    """
    if not caption_text:
        return 'g'  # Default to figure

    caption_lower = caption_text.lower()

    # Check for scheme indicators - be more specific
    if re.search(r'\bscheme\s+\d+', caption_lower) or re.search(r'\bscheme\s*:', caption_lower):
        return 's'

    # Check for figure indicators (including fig, figure)
    if re.search(r'\bfigure\s+\d+', caption_lower) or re.search(r'\bfig\s+\d+', caption_lower) or re.search(r'\bfig\.\s*\d+', caption_lower):
        return 'g'

    # If contains 'scheme' but not matched above, still consider it a scheme
    if 'scheme' in caption_lower:
        return 's'

    # Default to figure
    return 'g'


def extract_number_from_caption(caption_text: str, img_type: str) -> int:
    """
    Extract the number from a caption (e.g., "Figure 1" -> 1, "Scheme 2" -> 2).

    Args:
        caption_text: The caption text
        img_type: The image type ('g' for figure, 's' for scheme)

    Returns:
        The extracted number, or 0 if not found
    """
    if not caption_text:
        return 0

    caption_lower = caption_text.lower()

    # Search patterns based on type
    if img_type == 's':
        # Look for "Scheme 1", "Scheme: 1", "Scheme1", etc.
        patterns = [
            r'scheme\s*:?\s*(\d+)',
            r'scheme\s*(\d+)',
        ]
    else:
        # Look for "Figure 1", "Fig. 1", "Fig 1", "Figure: 1", etc.
        patterns = [
            r'figure\s*:?\s*(\d+)',
            r'fig\.\s*(\d+)',
            r'fig\s+(\d+)',
        ]

    for pattern in patterns:
        match = re.search(pattern, caption_lower)
        if match:
            return int(match.group(1))

    # If no number found, return 0
    return 0


def convert_wmf_emf_to_image(image_bytes: bytes, format_type: str) -> Image.Image:
    """
    Convert WMF/EMF file to PIL Image using platform-specific methods.

    Args:
        image_bytes: Raw bytes of the WMF/EMF file
        format_type: Format type (should be 'image/x-emf' or 'image/x-wmf')

    Returns:
        PIL Image object, or None if conversion fails

    Raises:
        Exception if conversion fails on all methods
    """
    is_windows = platform.system() == 'Windows'

    # Method 1: Try direct PIL conversion (works better on Windows)
    try:
        img = Image.open(BytesIO(image_bytes))

        if is_windows:
            # On Windows, PIL can use native GDI to render WMF/EMF
            # We need to explicitly load the image data
            img.load()

            # Convert to RGB
            if img.mode not in ('RGB', 'P'):
                img = img.convert('RGB')

            return img
        else:
            # On non-Windows, try to load and convert
            try:
                img.load()
                if img.mode not in ('RGB', 'P'):
                    img = img.convert('RGB')
                return img
            except:
                # If load fails, fall through to other methods
                pass
    except Exception as e:
        pass

    # Method 2: Try using Windows-specific pywin32 (only on Windows)
    if is_windows:
        try:
            import win32ui
            import win32con
            from PIL import ImageWin

            # Create a memory DC and load the EMF/WMF into it
            # This uses Windows GDI to render the metafile
            # Then capture it as a bitmap

            # Save bytes to temp file (required for Windows API)
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.emf' if 'emf' in format_type else '.wmf', delete=False) as tmp:
                tmp.write(image_bytes)
                tmp_path = tmp.name

            try:
                # Use Windows API to render the metafile
                import win32api

                # Get a DC
                desktop_dc = win32ui.CreateDCFromHandle(win32ui.GetDesktopWindow().GetDC().GetHandleOutput())

                # Create a compatible DC
                dc = desktop_dc.CreateCompatibleDC()

                # Play the metafile to get its dimensions and render it
                if 'emf' in format_type:
                    # Enhanced metafile
                    hemf = win32ui.CreateEnhMetaFile(dc, tmp_path)
                    # Render and convert to bitmap
                    # This is complex, so we'll use a simpler approach
                    pass

                # Simpler approach: Just try to open with PIL on Windows
                # which should work due to native support
                img = Image.open(tmp_path)
                img.load()

                if img.mode not in ('RGB', 'P'):
                    img = img.convert('RGB')

                os.unlink(tmp_path)
                return img

            except Exception as inner_e:
                try:
                    os.unlink(tmp_path)
                except:
                    pass
                raise inner_e

        except ImportError:
            # pywin32 not available
            pass
        except Exception as e:
            pass

    # Method 3: Try using wand (ImageMagick Python bindings)
    try:
        from wand.image import Image as WandImage

        with WandImage(blob=image_bytes) as wand_img:
            # Convert to PNG first for compatibility
            wand_img.format = 'png'
            png_blob = wand_img.make_blob()

            # Open with PIL
            img = Image.open(BytesIO(png_blob))
            return img

    except ImportError:
        # Wand not installed
        pass
    except Exception as e:
        pass

    # All methods failed
    raise Exception(f"Cannot convert WMF/EMF file: No suitable converter found. " +
                   f"On Windows, this should work natively. On Mac/Linux, install ImageMagick and wand: " +
                   f"pip install wand")


def resize_to_width(image: Image.Image, target_width: int = 1500) -> Image.Image:
    """
    Resize image to target width while maintaining aspect ratio.

    Args:
        image: PIL Image object
        target_width: Desired width in pixels

    Returns:
        Resized PIL Image object
    """
    original_width, original_height = image.size

    # Calculate proportional height
    aspect_ratio = original_height / original_width
    target_height = int(target_width * aspect_ratio)

    # Resize image using high-quality resampling
    resized_image = image.resize((target_width, target_height), Image.Resampling.LANCZOS)

    return resized_image


def convert_doc_to_docx(doc_path: str) -> str:
    """
    Convert .doc file to .docx format using available methods.

    Tries methods in order:
    1. win32com (Windows with MS Office) - Recommended
    2. doc2docx (if installed)
    3. Error if no method available

    Args:
        doc_path: Path to .doc file

    Returns:
        Path to converted .docx file

    Raises:
        RuntimeError: If no conversion method is available
    """
    output_path = doc_path.rsplit('.', 1)[0] + '_converted.docx'

    print(f"Converting {doc_path} to .docx format...")

    # Method 1: Try win32com (Windows with MS Office)
    if HAS_WIN32COM:
        try:
            import os
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            # Convert to absolute path
            abs_doc_path = os.path.abspath(doc_path)
            abs_output_path = os.path.abspath(output_path)

            doc = word.Documents.Open(abs_doc_path)
            doc.SaveAs(abs_output_path, FileFormat=16)  # 16 = wdFormatXMLDocument (.docx)
            doc.Close()
            word.Quit()

            print(f"  Converted using Microsoft Word COM")
            return output_path
        except Exception as e:
            print(f"  Warning: win32com conversion failed: {e}")
            # Fall through to next method

    # Method 2: Try doc2docx
    if HAS_DOC2DOCX:
        try:
            doc2docx.convert(doc_path, output_path)
            print(f"  Converted using doc2docx")
            return output_path
        except Exception as e:
            print(f"  Warning: doc2docx conversion failed: {e}")
            # Fall through to error

    # No method available
    raise RuntimeError(
        "Cannot convert .doc files: No conversion method available.\n"
        "Solutions:\n"
        "  1. Install pywin32: pip install pywin32 (requires MS Office on Windows)\n"
        "  2. Manually convert .doc to .docx in Microsoft Word\n"
        "  3. Install doc2docx: pip install pywin32 && pip install doc2docx --no-deps\n"
        f"\nHAS_WIN32COM={HAS_WIN32COM}, HAS_DOC2DOCX={HAS_DOC2DOCX}"
    )


def extract_images_from_docx(docx_path: str) -> List[Dict]:
    """
    Extract all images from a .docx file with their captions.
    Supports both modern (pic:pic) and VML (w:pict) image formats.

    Args:
        docx_path: Path to .docx file

    Returns:
        List of dictionaries containing image data and captions
    """
    doc = Document(docx_path)
    images = []

    # Track paragraph index for caption association
    for para_idx, paragraph in enumerate(doc.paragraphs):
        rel_ids = []

        # Check if paragraph contains modern format images (pic:pic)
        for run in paragraph.runs:
            if run._element.xpath('.//pic:pic'):
                # Found modern format image
                inline_shapes = run._element.xpath('.//a:blip/@r:embed')
                rel_ids.extend(inline_shapes)

        # Also check for VML format images (w:pict) - older Word format
        # These are typically in paragraphs, not necessarily in runs
        pict_elements = paragraph._element.xpath('.//w:pict')
        if pict_elements:
            # Look for v:imagedata elements which contain the relationship ID
            for pict in pict_elements:
                # Try multiple ways VML images might store the relationship ID
                # Get all elements and check their attributes
                all_elements = pict.xpath('.//*')
                for elem in all_elements:
                    # Check all attributes for relationship IDs
                    # VML images can use different attribute names
                    for attr_name, attr_value in elem.attrib.items():
                        # Look for attributes that contain 'rId' pattern (relationship IDs)
                        if attr_value and isinstance(attr_value, str) and attr_value.startswith('rId'):
                            rel_ids.append(attr_value)

        # Also check for embedded objects (w:object) - charts, diagrams, etc.
        # These contain images but are wrapped in object containers
        object_elements = paragraph._element.xpath('.//w:object')
        if object_elements:
            for obj in object_elements:
                # Look for v:imagedata elements within the object
                # OLE objects often have relationship IDs in v:imagedata or other child elements
                all_elements = obj.xpath('.//*')
                for elem in all_elements:
                    for attr_name, attr_value in elem.attrib.items():
                        if attr_value and isinstance(attr_value, str) and attr_value.startswith('rId'):
                            # Only add image relationship IDs, not embedding relationship IDs
                            # OLE objects have two rIds: one for the image preview, one for the .xls file
                            # We want the image preview
                            try:
                                if attr_value in doc.part.related_parts:
                                    part = doc.part.related_parts[attr_value]
                                    if part.content_type.startswith('image/'):
                                        rel_ids.append(attr_value)
                            except:
                                pass

        # Now process all found relationship IDs
        for rel_id in rel_ids:
            try:
                image_part = doc.part.related_parts[rel_id]

                # Skip non-image content types (Excel objects, etc.)
                content_type = image_part.content_type
                if not content_type.startswith('image/'):
                    continue

                image_bytes = image_part.blob

                # Search for caption in a wider range (3 paragraphs before and after)
                caption = ""
                search_range = 3

                # First check current paragraph
                caption = paragraph.text.strip()

                # If empty, search in nearby paragraphs
                if not caption:
                    for offset in range(1, search_range + 1):
                        # Check after
                        if para_idx + offset < len(doc.paragraphs):
                            next_text = doc.paragraphs[para_idx + offset].text.strip()
                            if next_text and ('figure' in next_text.lower() or 'scheme' in next_text.lower()):
                                caption = next_text
                                break

                        # Check before
                        if para_idx - offset >= 0:
                            prev_text = doc.paragraphs[para_idx - offset].text.strip()
                            if prev_text and ('figure' in prev_text.lower() or 'scheme' in prev_text.lower()):
                                caption = prev_text
                                break

                images.append({
                    'bytes': image_bytes,
                    'caption': caption,
                    'format': image_part.content_type,
                    'para_idx': para_idx  # Keep track of document order
                })
            except Exception as e:
                print(f"Warning: Could not extract image: {e}")
                continue

    # Remove duplicates based on image bytes
    unique_images = []
    seen_bytes = set()

    for img in images:
        img_hash = hash(img['bytes'])
        if img_hash not in seen_bytes:
            seen_bytes.add(img_hash)
            unique_images.append(img)

    return unique_images


def main(input_file: str, output_dir: str = None):
    """
    Main orchestration function.

    Args:
        input_file: Path to input Word document
        output_dir: Optional output directory (defaults to same as input file)
    """
    input_path = Path(input_file)

    if not input_path.exists():
        print(f"Error: File not found: {input_file}")
        sys.exit(1)

    # Determine output directory
    if output_dir:
        output_path = Path(output_dir) / input_path.stem
        output_path.mkdir(parents=True, exist_ok=True)
    else:
        # Create extracted_gifs subfolder with filename-based subfolder
        output_path = input_path.parent / "extracted_gifs" / input_path.stem
        output_path.mkdir(parents=True, exist_ok=True)

    # Extract base name for output files
    base_name = extract_document_basename(str(input_path))

    # Handle .doc vs .docx
    docx_file = str(input_path)
    temp_file = None

    if input_path.suffix.lower() == '.doc':
        temp_file = convert_doc_to_docx(str(input_path))
        docx_file = temp_file
    elif input_path.suffix.lower() != '.docx':
        print(f"Error: Unsupported file format: {input_path.suffix}")
        sys.exit(1)

    # Extract images
    print(f"Extracting images from {input_path.name}...")
    extracted_images = extract_images_from_docx(docx_file)

    if not extracted_images:
        print("No images found in document.")
        if temp_file and os.path.exists(temp_file):
            os.remove(temp_file)
        return

    print(f"Found {len(extracted_images)} image(s)")

    # Classify and extract numbers from all images first
    for img_data in extracted_images:
        img_type = classify_image(img_data['caption'])
        img_number = extract_number_from_caption(img_data['caption'], img_type)
        img_data['img_type'] = img_type
        img_data['img_number'] = img_number
        print(f"  Caption: '{img_data['caption'][:50]}...' -> Type: {img_type}, Number: {img_number}")

    # Sort images by document order (para_idx) to preserve original sequence
    # We'll still use separate counters for figures and schemes
    extracted_images.sort(key=lambda x: x.get('para_idx', 0))

    print(f"\nProcessing images in sorted order...")

    # Process and save images
    figure_counter = 1
    scheme_counter = 1
    figures_created = 0
    schemes_created = 0

    for img_data in extracted_images:
        try:
            # Check if this is a WMF/EMF file
            is_wmf_emf = img_data.get('format') in ('image/x-emf', 'image/x-wmf')

            if is_wmf_emf:
                # Use specialized WMF/EMF converter
                try:
                    image = convert_wmf_emf_to_image(img_data['bytes'], img_data.get('format'))
                except Exception as emf_error:
                    # If conversion fails, try fallback or skip
                    print(f"Warning: WMF/EMF conversion failed: {emf_error}")
                    print(f"  Caption: '{img_data.get('caption', '')[:50]}'")
                    print(f"  Format: {img_data.get('format')}")

                    # Save raw file as fallback
                    img_type = img_data['img_type']
                    if img_type == 's':
                        filename = f"{base_name}-s{scheme_counter:03d}.emf"
                        scheme_counter += 1
                    else:
                        filename = f"{base_name}-g{figure_counter:03d}.emf"
                        figure_counter += 1

                    output_file = output_path / filename
                    output_file.write_bytes(img_data['bytes'])
                    print(f"  Saved raw file: {filename} (needs manual conversion)")
                    continue
            else:
                # Load regular image from bytes
                image = Image.open(BytesIO(img_data['bytes']))

                # Handle WMF/EMF files detected by PIL
                if image.format in ('WMF', 'EMF'):
                    try:
                        image = convert_wmf_emf_to_image(img_data['bytes'], img_data.get('format'))
                    except Exception as emf_error:
                        print(f"Warning: WMF/EMF file detected but conversion failed: {emf_error}")
                        # Fall through to try regular processing
                        pass

            # Convert to RGB if necessary (for GIF compatibility)
            if image.mode not in ('RGB', 'P'):
                image = image.convert('RGB')

            # Resize to target width
            resized_image = resize_to_width(image, target_width=1500)

            # Use pre-classified image type
            img_type = img_data['img_type']

            # Generate filename
            if img_type == 's':
                filename = f"{base_name}-s{scheme_counter:03d}.gif"
                scheme_counter += 1
                schemes_created += 1
            else:
                filename = f"{base_name}-g{figure_counter:03d}.gif"
                figure_counter += 1
                figures_created += 1

            # Save as GIF
            output_file = output_path / filename
            resized_image.save(output_file, 'GIF', optimize=True)

            width, height = resized_image.size
            print(f"Created: {filename} ({width}x{height})")

        except Exception as e:
            print(f"Error processing image (caption: '{img_data.get('caption', '')[:50]}'): {e}")
            continue

    # Clean up temp file
    if temp_file and os.path.exists(temp_file):
        os.remove(temp_file)

    # Summary
    print(f"\nSummary: {figures_created} figure(s), {schemes_created} scheme(s) extracted")
    print(f"Output directory: {output_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Extract images from Word documents and convert to GIF format"
    )
    parser.add_argument(
        "input_file",
        help="Path to input Word document (.doc or .docx)"
    )
    parser.add_argument(
        "-o", "--output",
        help="Output directory (default: extracted_gifs/ in same directory as input)",
        default=None
    )

    args = parser.parse_args()
    main(args.input_file, args.output)
