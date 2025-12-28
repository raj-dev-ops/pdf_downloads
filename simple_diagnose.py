#!/usr/bin/env python3
"""
Simple diagnostic to find image containers
"""

import sys
from docx import Document

def simple_diagnose(docx_path):
    """Find where images are stored"""

    doc = Document(docx_path)

    print(f"Analyzing: {docx_path}\n")

    # Get all image relationships
    image_rels = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            image_rels[rel_id] = rel.target_ref

    print(f"Found {len(image_rels)} embedded images\n")

    # Check for w:pict (Picture elements - VML format, older Word compatibility)
    pict_count = 0
    for para in doc.paragraphs:
        picts = para._element.xpath('.//w:pict')
        if picts:
            pict_count += len(picts)
            print(f"Found {len(picts)} w:pict elements in a paragraph")
            # Check what's inside
            for pict in picts:
                # Look for v:shape elements
                shapes = pict.xpath('.//v:shape')
                if shapes:
                    print(f"  Contains {len(shapes)} v:shape element(s)")
                # Look for v:imagedata
                imagedatas = pict.xpath('.//v:imagedata')
                if imagedatas:
                    print(f"  Contains {len(imagedatas)} v:imagedata element(s)")
                    for imagedata in imagedatas:
                        # Get all attributes
                        attrs = imagedata.attrib
                        print(f"    Attributes: {attrs}")

    print(f"\nTotal w:pict elements: {pict_count}")

    # Check for mc:AlternateContent (compatibility wrappers)
    alt_count = 0
    for para in doc.paragraphs:
        alts = para._element.xpath('.//mc:AlternateContent')
        if alts:
            alt_count += len(alts)
            print(f"\nFound {len(alts)} mc:AlternateContent in a paragraph")
            for alt in alts:
                # Check what's inside Choice and Fallback
                choices = alt.xpath('.//mc:Choice')
                fallbacks = alt.xpath('.//mc:Fallback')
                print(f"  Contains {len(choices)} Choice(s) and {len(fallbacks)} Fallback(s)")

                # Look for drawings in Choice
                for choice in choices:
                    drawings = choice.xpath('.//w:drawing')
                    print(f"    Choice has {len(drawings)} drawing(s)")

                # Look for pict in Fallback
                for fallback in fallbacks:
                    picts = fallback.xpath('.//w:pict')
                    print(f"    Fallback has {len(picts)} w:pict(s)")

    print(f"\nTotal mc:AlternateContent elements: {alt_count}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python simple_diagnose.py <path_to_docx>")
        sys.exit(1)

    simple_diagnose(sys.argv[1])
