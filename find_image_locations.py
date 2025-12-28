#!/usr/bin/env python3
"""
Find where specific images are located in the document
"""

import sys
from docx import Document

def find_image_in_tables(doc, target_rel_id):
    """Search for image in tables"""
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                # Check in cell paragraphs
                for para in cell.paragraphs:
                    # Check modern format
                    for run in para.runs:
                        if run._element.xpath('.//pic:pic'):
                            inline_shapes = run._element.xpath('.//a:blip/@r:embed')
                            if target_rel_id in inline_shapes:
                                return f"Table {table_idx}, Row {row_idx}, Cell {cell_idx}"

                    # Check VML format
                    pict_elements = para._element.xpath('.//w:pict')
                    for pict in pict_elements:
                        all_elements = pict.xpath('.//*')
                        for elem in all_elements:
                            for attr_name, attr_value in elem.attrib.items():
                                if attr_value == target_rel_id:
                                    return f"Table {table_idx}, Row {row_idx}, Cell {cell_idx}"
    return None

def find_image_in_sections(doc, target_rel_id):
    """Search for image in headers/footers"""
    for section_idx, section in enumerate(doc.sections):
        # Check header
        try:
            header = section.header
            for para in header.paragraphs:
                # Check modern format
                for run in para.runs:
                    if run._element.xpath('.//pic:pic'):
                        inline_shapes = run._element.xpath('.//a:blip/@r:embed')
                        if target_rel_id in inline_shapes:
                            return f"Section {section_idx} Header"

                # Check VML format
                pict_elements = para._element.xpath('.//w:pict')
                for pict in pict_elements:
                    all_elements = pict.xpath('.//*')
                    for elem in all_elements:
                        for attr_name, attr_value in elem.attrib.items():
                            if attr_value == target_rel_id:
                                return f"Section {section_idx} Header"
        except:
            pass

        # Check footer
        try:
            footer = section.footer
            for para in footer.paragraphs:
                # Check modern format
                for run in para.runs:
                    if run._element.xpath('.//pic:pic'):
                        inline_shapes = run._element.xpath('.//a:blip/@r:embed')
                        if target_rel_id in inline_shapes:
                            return f"Section {section_idx} Footer"

                # Check VML format
                pict_elements = para._element.xpath('.//w:pict')
                for pict in pict_elements:
                    all_elements = pict.xpath('.//*')
                    for elem in all_elements:
                        for attr_name, attr_value in elem.attrib.items():
                            if attr_value == target_rel_id:
                                return f"Section {section_idx} Footer"
        except:
            pass

    return None

def analyze_missing_images(docx_path):
    """Find where missing images are located"""

    doc = Document(docx_path)

    missing_ids = ['rId20', 'rId7']

    print("Searching for missing images:")
    print("=" * 70)

    for rel_id in missing_ids:
        rel = doc.part.rels.get(rel_id)
        if rel:
            print(f"\n{rel_id} -> {rel.target_ref}")

            # Search in tables
            location = find_image_in_tables(doc, rel_id)
            if location:
                print(f"  Found in: {location}")
                continue

            # Search in headers/footers
            location = find_image_in_sections(doc, rel_id)
            if location:
                print(f"  Found in: {location}")
                continue

            print(f"  NOT FOUND in paragraphs, tables, headers, or footers")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python find_image_locations.py <path_to_docx>")
        sys.exit(1)

    analyze_missing_images(sys.argv[1])
