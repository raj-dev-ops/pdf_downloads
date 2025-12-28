#!/usr/bin/env python3
"""
Check which images are being deduplicated
"""

import sys
from docx import Document
from io import BytesIO
from PIL import Image

def check_duplicates(docx_path):
    """Check duplicate image detection"""

    doc = Document(docx_path)
    images_data = []

    # Extract all images
    for para_idx, paragraph in enumerate(doc.paragraphs):
        rel_ids = []

        # Modern format
        for run in paragraph.runs:
            if run._element.xpath('.//pic:pic'):
                inline_shapes = run._element.xpath('.//a:blip/@r:embed')
                rel_ids.extend(inline_shapes)

        # VML format
        pict_elements = paragraph._element.xpath('.//w:pict')
        for pict in pict_elements:
            all_elements = pict.xpath('.//*')
            for elem in all_elements:
                for attr_name, attr_value in elem.attrib.items():
                    if attr_value and isinstance(attr_value, str) and attr_value.startswith('rId'):
                        rel_ids.append(attr_value)

        # Embedded objects
        object_elements = paragraph._element.xpath('.//w:object')
        for obj in object_elements:
            all_elements = obj.xpath('.//*')
            for elem in all_elements:
                for attr_name, attr_value in elem.attrib.items():
                    if attr_value and isinstance(attr_value, str) and attr_value.startswith('rId'):
                        rel_ids.append(attr_value)

        # Process
        for rel_id in rel_ids:
            try:
                image_part = doc.part.related_parts[rel_id]
                image_bytes = image_part.blob
                caption = paragraph.text.strip()

                # Search nearby for caption
                if not caption:
                    for offset in range(1, 4):
                        if para_idx + offset < len(doc.paragraphs):
                            next_text = doc.paragraphs[para_idx + offset].text.strip()
                            if next_text and ('figure' in next_text.lower() or 'scheme' in next_text.lower()):
                                caption = next_text
                                break
                        if para_idx - offset >= 0:
                            prev_text = doc.paragraphs[para_idx - offset].text.strip()
                            if prev_text and ('figure' in prev_text.lower() or 'scheme' in prev_text.lower()):
                                caption = prev_text
                                break

                images_data.append({
                    'rel_id': rel_id,
                    'bytes_len': len(image_bytes),
                    'bytes_hash': hash(image_bytes),
                    'caption': caption[:80] if caption else '(no caption)',
                    'para_idx': para_idx,
                    'format': image_part.content_type
                })

            except Exception as e:
                print(f"Error extracting {rel_id}: {e}")

    print(f"Total images found: {len(images_data)}")
    print("\n" + "=" * 80)
    print("All images in order:")
    print("=" * 80)

    for idx, img in enumerate(images_data):
        print(f"{idx+1}. {img['rel_id']} (para {img['para_idx']}) - {img['format']}")
        print(f"   Caption: {img['caption']}")
        print(f"   Bytes: {img['bytes_len']}, Hash: {img['bytes_hash']}")

    # Check for duplicates
    print("\n" + "=" * 80)
    print("Checking for duplicate byte hashes:")
    print("=" * 80)

    seen_hashes = {}
    for idx, img in enumerate(images_data):
        h = img['bytes_hash']
        if h in seen_hashes:
            print(f"\nDUPLICATE FOUND:")
            print(f"  Original: #{seen_hashes[h]+1} - {images_data[seen_hashes[h]]['caption']}")
            print(f"  Duplicate: #{idx+1} - {img['caption']}")
        else:
            seen_hashes[h] = idx

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python check_duplicates.py <path_to_docx>")
        sys.exit(1)

    check_duplicates(sys.argv[1])
